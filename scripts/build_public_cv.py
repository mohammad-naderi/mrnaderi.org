from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "public" / "documents" / "Mohammad-Reza-Naderi-CV.docx"

# compact_reference_guide preset, with one named site-identity color override.
FONT = "Calibri"
INK = RGBColor(29, 36, 33)
BLUE = RGBColor(56, 87, 131)
MUTED = RGBColor(97, 104, 98)
LINE = "CBC8BB"


def set_font(run, size=None, bold=None, italic=None, color=None, caps=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if caps is not None:
        run.font.all_caps = caps


def set_cell_borderless(_):
    return


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    lead = doc.styles.add_style("CV Lead", WD_STYLE_TYPE.PARAGRAPH)
    lead.base_style = normal
    lead.font.name = FONT
    lead._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    lead._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    lead.font.size = Pt(11.5)
    lead.font.color.rgb = INK
    lead.paragraph_format.space_after = Pt(12)
    lead.paragraph_format.line_spacing = 1.25

    entry = doc.styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
    entry.base_style = normal
    entry.font.name = FONT
    entry._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    entry._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    entry.font.size = Pt(10.5)
    entry.paragraph_format.left_indent = Inches(0.82)
    entry.paragraph_format.first_line_indent = Inches(-0.82)
    entry.paragraph_format.tab_stops.add_tab_stop(Inches(0.82))
    entry.paragraph_format.space_after = Pt(5)
    entry.paragraph_format.line_spacing = 1.2
    entry.paragraph_format.widow_control = True


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)
    p_pr.append(borders)


def add_entry(doc, year, title, details="", title_italic=False):
    p = doc.add_paragraph(style="CV Entry")
    year_run = p.add_run(year)
    set_font(year_run, size=9.5, bold=True, color=BLUE)
    p.add_run("\t")
    title_run = p.add_run(title)
    set_font(title_run, size=10.5, bold=not title_italic, italic=title_italic, color=INK)
    if details:
        detail_run = p.add_run(details)
        set_font(detail_run, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    hrun = header.add_run("CURRICULUM VITAE  /  AUGUST 2026")
    set_font(hrun, size=8.5, bold=True, color=MUTED, caps=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    frun = footer.add_run("MOHAMMAD REZA NADERI  |  PAGE ")
    set_font(frun, size=8.5, bold=True, color=MUTED, caps=True)
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(14)
    kicker.paragraph_format.space_after = Pt(4)
    krun = kicker.add_run("PHILOSOPHER  /  AUTHOR  /  TEACHER")
    set_font(krun, size=9, bold=True, color=BLUE, caps=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    trun = title.add_run("Mohammad Reza Naderi")
    set_font(trun, size=28, bold=False, color=INK)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    mrun = meta.add_run("Toronto, Canada  |  mrnaderi.org")
    set_font(mrun, size=10.5, color=MUTED)
    add_rule(doc)

    profile = doc.add_paragraph(style="CV Lead")
    profile.add_run(
        "Philosopher working across continental philosophy, mathematical ontology, "
        "psychoanalysis, and political thought. His Theory of Discipline investigates "
        "how truths begin, acquire consistency, and become universal and absolute. "
        "His current work places radical novelty at the centre of philosophy."
    )

    add_heading(doc, "Research", 1)
    add_entry(
        doc,
        "Areas",
        "Alain Badiou; truth and radical novelty; mathematical ontology and logic; "
        "theory of the subject; Hegel and Lacan; philosophy of disciplines.",
    )
    add_entry(
        doc,
        "Current",
        "A third book on the contemporary foreclosure of novelty and the philosophical "
        "conditions of radical novelty; ongoing development of the Theory of Discipline.",
    )

    add_heading(doc, "Education", 1)
    add_entry(
        doc,
        "2017",
        "PhD, Philosophy and Critical Thinking, European Graduate School, summa cum laude.",
        " Dissertation: Infinity and Subjectivity: Badiou's Theory of Discipline. "
        "Supervisor: Alain Badiou.",
    )
    add_entry(
        doc,
        "2006",
        "MA, Philosophy, University of Toronto.",
        " Research focus: Heidegger and German Idealism.",
    )
    add_entry(
        doc,
        "1994",
        "BSc, Computer Science and Pure Mathematics, University of Toronto.",
    )

    add_heading(doc, "Books", 1)
    add_entry(
        doc,
        "2027",
        "Badiou and Disciplinary Truths: A Companion to The Immanence of Truths.",
        " Bloomsbury Academic. Forthcoming.",
        title_italic=True,
    )
    add_entry(
        doc,
        "2024",
        "Philosophy After Lacan: Politics, Science, and Art.",
        " Co-edited with Alireza Taheri and Chris Vanderwees. Routledge.",
        title_italic=True,
    )
    add_entry(
        doc,
        "2023",
        "Badiou, Infinity, and Subjectivity: Reading Hegel and Lacan after Badiou.",
        " Lexington Books.",
        title_italic=True,
    )

    add_heading(doc, "Articles, Chapters, and Collective Works", 1)
    entries = [
        (
            "2024",
            '"Love - The Scene of the Two."',
            " Poliética 12(4): 164-190.",
        ),
        (
            "2024",
            '"Lacan Is Our Hegel: Dialectic from Hegel to Lacan to Badiou."',
            " In Philosophy After Lacan: Politics, Science, and Art, 142-166. Routledge.",
        ),
        (
            "2023",
            '"The Place of the Subject in Badiou\'s Theory of Discipline."',
            " Filozofski vestnik 43(3).",
        ),
        (
            "2021",
            '"Atlas of Experimental Politics."',
            " Collective work, including 'The Discipline of Politics.' ŠUM 17: Meta-Futures.",
        ),
        (
            "2020",
            '"Contribution to the Critique of Political Organization."',
            " Collective work. Crisis & Critique 7(3): 398-428.",
        ),
        (
            "2020",
            '"Model Theory and the Unnamable."',
            " In Revolutions for the Future: May '68 and the Prague Spring, 92-121. Suture Press.",
        ),
        (
            "2018",
            '"Mark and Lack: Formalism as Fidelity."',
            " Crisis & Critique 5(1): 272-299.",
        ),
        (
            "2008",
            '"War, Peace and Fuzzy Logic."',
            " With Mory Ghomshei and John Meech. Cybernetics and Systems 39(2): 113-135.",
        ),
        (
            "2008",
            '"Fuzzy Logic in a Postmodern Era."',
            " With Mory Ghomshei and John Meech. In Forging New Frontiers: Fuzzy Pioneers II, 363-376.",
        ),
    ]
    for year, title_text, details in entries:
        add_entry(doc, year, title_text, details)

    add_heading(doc, "Teaching and Presentations", 1)
    add_entry(
        doc,
        "2026",
        "Adventures of French Structuralism.",
        " Seminar sequence, Toronto Psychoanalytic Society & Institute: "
        '"Meaning, Structure, and the Problem of Determination" (16 April); '
        '"From Structuralism to Process: Deleuze, Anti-Oedipus, and the Problem of Novelty" (22 April).',
        title_italic=True,
    )
    add_entry(
        doc,
        "2026",
        '"Dark Enlightenment."',
        " Two-session invited presentation, 8 and 22 May.",
    )
    add_entry(
        doc,
        "2019",
        '"The Unconscious Materialism of Hegel."',
        " Dialectic Returns, Prague, September.",
    )

    add_heading(doc, "Affiliations", 1)
    add_entry(
        doc,
        "Member",
        "Subset of Theoretical Practice.",
        " Research collective working across logic, linguistics, psychoanalysis, and political theory.",
    )
    add_entry(
        doc,
        "Teacher",
        "Toronto Psychoanalytic Society & Institute.",
        " Contemporary French philosophy.",
    )

    add_heading(doc, "Selected Professional Experience", 1)
    add_entry(
        doc,
        "2005-",
        "Founder and principal, FARA Logics Inc., Toronto.",
        " Enterprise architecture, software systems, payments, data governance, and technology advisory work "
        "for major Canadian and American institutions.",
    )
    add_entry(
        doc,
        "Earlier",
        "Software architect, entrepreneur, consultant, and technology leader.",
        " Long-term work included CIBC; other engagements included BEA Systems, General Motors, "
        "Wells Fargo, Kmart, HP, IBM, and Manulife.",
    )

    core = doc.core_properties
    core.title = "Curriculum Vitae - Mohammad Reza Naderi"
    core.subject = "Public intellectual and academic curriculum vitae"
    core.author = "Mohammad Reza Naderi"
    core.keywords = "philosophy, Badiou, Theory of Discipline, curriculum vitae"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
